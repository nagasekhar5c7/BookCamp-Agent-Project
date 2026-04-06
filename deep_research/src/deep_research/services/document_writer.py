"""Word document writer — renders a validated outline to .docx.

Deterministic, no LLM calls. Consumes a :class:`ReportOutline` and a
:class:`CitationRegistry` and produces a plain-but-clean ``.docx`` with:

1. Title heading.
2. Each outline section rendered as ``Heading 1`` + its paragraphs
   verbatim. Inline citation markers (``[1]``, ``[2][5]``) are left as
   literal text — no hyperlinking in v1.
3. A final **References** section listing every citation in id order
   as ``[n] Title — URL (accessed YYYY-MM-DD)``.

A pre-render validation pass rejects any paragraph containing a
citation marker that isn't registered. This is the enforcement point
for ideas.md §6 step 5 ("orphan markers abort the job with a clear
error").
"""

from __future__ import annotations

import re
from pathlib import Path

import structlog
from docx import Document

from deep_research.models.outline import ReportOutline
from deep_research.services.citation_registry import CitationRegistry

log = structlog.get_logger(__name__)

# Matches one inline marker like "[3]". The surrounding "[1][2]" style
# is just two adjacent matches, which the finditer loop handles naturally.
_CITATION_MARKER_RE = re.compile(r"\[(\d+)\]")


class OrphanCitationError(ValueError):
    """Raised when a paragraph cites an id not present in the registry."""


def write_document(
    *,
    outline: ReportOutline,
    registry: CitationRegistry,
    output_dir: Path | str,
    job_id: str,
) -> Path:
    """Render ``outline`` to a .docx file and return its path.

    Args:
        outline: The validated report outline from the Lead Synthesizer.
        registry: The citation registry — every inline ``[n]`` marker in
            the outline must resolve to an id registered here.
        output_dir: Directory to write the file into. Created if missing.
        job_id: Used in the generated filename so multiple jobs never
            collide on disk.

    Returns:
        Absolute path to the written ``.docx``.

    Raises:
        OrphanCitationError: If any paragraph references a citation id
            that is not in the registry.
    """
    bound_log = log.bind(job_id=job_id)

    _validate_citations(outline, registry, log=bound_log)

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    path = (output_dir / f"research_{job_id}.docx").resolve()

    doc = Document()
    doc.add_heading(outline.title, level=0)

    for section in outline.sections:
        doc.add_heading(section.heading, level=1)
        for paragraph in section.paragraphs:
            doc.add_paragraph(paragraph)

    # References section — always present, even if empty, so the reader
    # can see immediately whether any sources were used.
    doc.add_heading("References", level=1)
    if len(registry) == 0:
        doc.add_paragraph("(no sources were cited)")
    else:
        for cid, citation in registry.items():
            accessed = citation.accessed_at.strftime("%Y-%m-%d")
            doc.add_paragraph(
                f"[{cid}] {citation.title} — {citation.url} (accessed {accessed})"
            )

    doc.save(str(path))
    bound_log.info("document_written", path=str(path), sections=len(outline.sections))
    return path


# --------------------------------------------------------------------------- #
# Internals                                                                   #
# --------------------------------------------------------------------------- #


def _validate_citations(
    outline: ReportOutline,
    registry: CitationRegistry,
    *,
    log: structlog.stdlib.BoundLogger,
) -> None:
    """Scan every paragraph for orphan citation markers.

    Walks all paragraphs once, collects every ``[n]`` id, and checks
    each against the registry. A single orphan id raises — partial
    documents never leave the system.
    """
    orphan_ids: set[int] = set()
    for section in outline.sections:
        for paragraph in section.paragraphs:
            for match in _CITATION_MARKER_RE.finditer(paragraph):
                cid = int(match.group(1))
                if not registry.contains(cid):
                    orphan_ids.add(cid)

    if orphan_ids:
        log.error("orphan_citation_markers", orphan_ids=sorted(orphan_ids))
        raise OrphanCitationError(
            f"Outline references citation ids not in the registry: "
            f"{sorted(orphan_ids)}"
        )
